from docx.api import Document
import json, sys

if __name__=="__main__":
    coversheet_data = json.loads(sys.argv[1])
    document = Document('py-docx/Monitoring-Cover-Sheet.docx')
    i=0
    j=0
    dictionary = coversheet_data['contact_details']

    standard_criteria = coversheet_data['monitoring_criteria']
    criteria = standard_criteria['standard_criteria'].split('#')
    p_criteria = standard_criteria['project_criteria'].split('#')

    for rows in document.tables[0].rows:
        head = rows.cells[0].text
        if head in dictionary:
            if head in ['Protocol Title :', 'Monitoring Start Date :']:
                rows.cells[1].paragraphs[0].add_run(dictionary[head]).bold = True
            else: 
                if 'Monitor 3' in head:
                    if i==0:
                        rows.cells[1].paragraphs[0].add_run(dictionary[head][0]).bold = True
                        rows.cells[2].paragraphs[0].add_run(dictionary[head][1]).bold = True
                    else: 
                        rows.cells[1].paragraphs[0].add_run(" - Ticked " if  dictionary['Supervisor :'] else " - Unticked").bold = True
                    i+=1
                else:
                    rows.cells[1].paragraphs[0].add_run(dictionary[head][0]).bold = True
                    rows.cells[2].paragraphs[0].add_run(dictionary[head][1]).bold = True
    document.tables[1].rows[0].cells[0].paragraphs[0].add_run(coversheet_data['species_phenotype_issues']['Species']).bold = True

    #check there are how many criteria and add rows to the tables if need
    if len(criteria) > 4:
        for i in range(len(criteria)-4):
            document.tables[2].add_row()
    
    x = 0 #to access the element of criteria
    jj = 0 # used to skiped the first row
    for rows in document.tables[2].rows:
        if jj > 0 and jj <= len(criteria):
            temp = criteria[x].split('@')
            x += 1
            for k in range(4):
                rows.cells[k].paragraphs[0].add_run(temp[k]).bold = True
        jj += 1

    #check there are how many criteria and add rows to the tables if need
    if len(p_criteria) > 2:
        for i in range(len(p_criteria)-2):
            document.tables[3].add_row()

    x = 0 #to access the element of criteria
    jj = 0 # used to skiped the first row
    for rows in document.tables[3].rows:
        if jj > 0 and jj <= len(p_criteria):
            temp = p_criteria[x].split('@')
            x += 1
            for k in range(4):
                rows.cells[k].paragraphs[0].add_run(temp[k]).bold = True
        jj += 1

    #MONITORING FREQUENCY
    document.tables[5].rows[0].cells[0].paragraphs[0].add_run(coversheet_data['monitoring_frequency']['monitoring_frequency']).bold = True

    index = 0
    type_of_recordingsheet = coversheet_data['type_of_recording_sheet']
    if type_of_recordingsheet['general'] == True:
        index = 0
    elif type_of_recordingsheet['anasthesia'] == True:
        index = 1
    elif type_of_recordingsheet['post_proc'] == True:
        index = 2
    elif type_of_recordingsheet['other'] == True:
        index = 3
        other_D = type_of_recordingsheet['other_description']

    if index != 3:
        document.tables[6].rows[1].cells[index].paragraphs[0].add_run("X").bold = True
    else:
        document.tables[6].rows[1].cells[index].paragraphs[0].add_run(other_D).bold = True

    document.save('animalwellbeing/static/animalwellbeing/coversheets/{}.docx'.format(sys.argv[2]))

